import os
from io import StringIO

import matplotlib.lines as lines
import matplotlib.patches as patches
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from Bio import pairwise2, SeqIO
from Bio.Blast import NCBIXML
from Bio.Blast.Applications import NcbiblastpCommandline
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord
from Bio.SubsMat import MatrixInfo as matlist
from PyQt5 import QtWidgets
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import Pt
from modlamp.core import load_scale
from modlamp.descriptors import PeptideDescriptor


class Helper:
    data_filename = 'gpcr_data.xlsx'
    summary_filename = 'result-summary.docx'
    df_template = pd.read_excel(data_filename, sheet_name="Templates")
    df_target = pd.read_excel(data_filename, sheet_name="Targets")
    df_pdb = pd.read_excel(data_filename, sheet_name="Templates", column_name="Sequence")
    df_active = pd.read_excel('gpcr_data.xlsx', sheet_name="active")
    df_inactive = pd.read_excel('gpcr_data.xlsx', sheet_name="inactive")
    df_intermediate = pd.read_excel('gpcr_data.xlsx', sheet_name="intermediate")

    images_dir = 'images'
    valid_residues = {'A', 'R', 'N', 'D', 'C', 'Q', 'E', 'G', 'H', 'I', 'L', 'K', 'M', 'F', 'P', 'S', 'T', 'W', 'Y', 'V'}

    # eisenberg
    hydscale = {'A': 0.620, 'R': -2.530, 'N': -0.780, 'D': -0.900, 'C': 0.290,
                'Q': -0.850, 'E': -0.740, 'G': 0.480, 'H': -0.400, 'I': 1.380,
                'L': 1.060, 'K': -1.500, 'M': 0.640, 'F': 1.190, 'P': 0.120,
                'S': -0.180, 'T': -0.050, 'W': 0.810, 'Y': 0.260, 'V': 1.080}

    non_olfactory_hotspot_residue_positions = [3.28, 3.29, 3.33, 3.36, 3.37, 5.39, 5.40, 5.43, 5.44, 5.47,
                                               6.44, 6.51, 6.52, 6.55, 6.58, 7.31, 7.34, 7.38, 7.41, 7.42,
                                               2.60, 2.63, 3.32, 5.53]
    gpcr_tm = {
        ("A", "A"): 2, ("A", "C"): 0, ("A", "P"): -1, ("A", "S"): 1,
        ("A", "T"): 1, ("C", "C"): 5, ("D", "A"): -3, ("D", "C"): -3,
        ("D", "D"): 9, ("D", "G"): -1, ("D", "N"): 2, ("D", "P"): 0,
        ("D", "S"): -2, ("D", "T"): -2, ("E", "A"): -1, ("E", "C"): -3,
        ("E", "D"): 5, ("E", "E"): 6, ("E", "G"): 1, ("E", "N"): 0,
        ("E", "P"): 0, ("E", "S"): 1, ("E", "T"): -1, ("F", "A"): -1,
        ("F", "C"): 0, ("F", "D"): -2, ("F", "E"): -2, ("F", "F"): 2,
        ("F", "G"): 0, ("F", "H"): -1, ("F", "I"): 0, ("F", "K"): -1,
        ("F", "L"): 0, ("F", "M"): 0, ("F", "N"): -2, ("F", "P"): -3,
        ("F", "Q"): -1, ("F", "R"): -2, ("F", "S"): -1, ("F", "T"): -1,
        ("F", "V"): 0, ("G", "A"): 1, ("G", "C"): 0, ("G", "G"): 4,
        ("G", "P"): -2, ("G", "S"): 1, ("G", "T"): 1, ("H", "A"): -1,
        ("H", "C"): -2, ("H", "D"): 1, ("H", "E"): 1, ("H", "G"): -1,
        ("H", "H"): 5, ("H", "N"): 2, ("H", "P"): 0, ("H", "Q"): 3,
        ("H", "S"): 0, ("H", "T"): 0, ("I", "A"): -1, ("I", "C"): -1,
        ("I", "D"): -4, ("I", "E"): -2, ("I", "G"): -1, ("I", "H"): -1,
        ("I", "I"): 2, ("I", "K"): -2, ("I", "M"): 1, ("I", "N"): -3,
        ("I", "P"): -3, ("I", "Q"): -2, ("I", "R"): -3, ("I", "S"): -1,
        ("I", "T"): 0, ("K", "A"): -1, ("K", "C"): -2, ("K", "D"): 0,
        ("K", "E"): 3, ("K", "G"): -1, ("K", "H"): 2, ("K", "K"): 6,
        ("K", "N"): 0, ("K", "P"): -3, ("K", "Q"): 3, ("K", "R"): 4,
        ("K", "S"): 0, ("K", "T"): -1, ("L", "A"): -1, ("L", "C"): -1,
        ("L", "D"): -4, ("L", "E"): -2, ("L", "G"): -1, ("L", "H"): -1,
        ("L", "I"): 1, ("L", "K"): -2, ("L", "L"): 2, ("L", "M"): 1,
        ("L", "N"): -3, ("L", "P"): -3, ("L", "Q"): -1, ("L", "R"): -2,
        ("L", "S"): -1, ("L", "T"): 0, ("M", "A"): -1, ("M", "C"): 0,
        ("M", "D"): -3, ("M", "E"): -1, ("M", "G"): 0, ("M", "H"): 0,
        ("M", "K"): -1, ("M", "M"): 3, ("M", "N"): -2, ("M", "P"): -3,
        ("M", "Q"): -1, ("M", "R"): -2, ("M", "S"): -1, ("M", "T"): 0,
        ("N", "A"): -2, ("N", "C"): -3, ("N", "G"): -2, ("N", "N"): 8,
        ("N", "P"): -2, ("N", "S"): 0, ("N", "T"): -1, ("P", "C"): -3,
        ("P", "P"): 8, ("P", "S"): 0, ("P", "T"): -1, ("Q", "A"): -1,
        ("Q", "C"): -2, ("Q", "D"): 1, ("Q", "E"): 3, ("Q", "G"): -1,
        ("Q", "N"): 1, ("Q", "P"): 0, ("Q", "Q"): 5, ("Q", "S"): 0,
        ("Q", "T"): -1, ("R", "A"): -1, ("R", "C"): -1, ("R", "D"): -1,
        ("R", "E"): 2, ("R", "G"): -1, ("R", "H"): 2, ("R", "N"): 0,
        ("R", "P"): -2, ("R", "Q"): 3, ("R", "R"): 7, ("R", "S"): -1,
        ("R", "T"): -1, ("S", "C"): 0, ("S", "S"): 2, ("T", "C"): 0,
        ("T", "S"): 1, ("T", "T"): 2, ("V", "A"): 0, ("V", "C"): 0,
        ("V", "D"): -3, ("V", "E"): -2, ("V", "G"): -1, ("V", "H"): -1,
        ("V", "I"): 1, ("V", "K"): -2, ("V", "L"): 0, ("V", "M"): 0,
        ("V", "N"): -3, ("V", "P"): -3, ("V", "Q"): -1, ("V", "R"): -2,
        ("V", "S"): -1, ("V", "T"): 0, ("V", "V"): 1, ("W", "A"): -2,
        ("W", "C"): -1, ("W", "D"): -3, ("W", "E"): -1, ("W", "F"): 0,
        ("W", "G"): -1, ("W", "H"): 0, ("W", "I"): -2, ("W", "K"): -1,
        ("W", "L"): -1, ("W", "M"): -1, ("W", "N"): -2, ("W", "P"): -4,
        ("W", "Q"): 1, ("W", "R"): 0, ("W", "S"): -2, ("W", "T"): -2,
        ("W", "V"): -1, ("W", "W"): 8, ("W", "Y"): 2, ("Y", "A"): -2,
        ("Y", "C"): -1, ("Y", "D"): -2, ("Y", "E"): -1, ("Y", "F"): 1,
        ("Y", "G"): -1, ("Y", "H"): 1, ("Y", "I"): -2, ("Y", "K"): -1,
        ("Y", "L"): -1, ("Y", "M"): -1, ("Y", "N"): -1, ("Y", "P"): -4,
        ("Y", "Q"): 0, ("Y", "R"): -2, ("Y", "S"): -2, ("Y", "T"): -2,
        ("Y", "V"): -1, ("Y", "Y"): 6
    }

    @staticmethod
    def calculate_hotspot_residue_similarity(targ_seq, temp_seq, hotspot_residue_position):
        matrix = Helper.gpcr_tm

        targ_char = Helper._fetch_sequence_character(targ_seq, temp_seq, 'target', hotspot_residue_position)
        temp_char = Helper._fetch_sequence_character(targ_seq, temp_seq, 'template', hotspot_residue_position)

        char_pair = (targ_char, temp_char)
        char_pair_reverse = (temp_char, targ_char)
        value = matrix[char_pair] if char_pair in matrix else matrix[char_pair_reverse]

        return value

    @staticmethod
    def _fetch_sequence_character(targ_seq, temp_seq, target_or_template, hotspot_residue_position):
        # hotspot_residue_position = 3.28
        reference_number = int(hotspot_residue_position)  # 3
        bw_value = ''

        if target_or_template == 'target':
            bw_value = Helper.df_target.loc[Helper.df_target['Sequence'] == targ_seq, 'BW' + str(reference_number) + '.50'].iloc[0]  # R132
        elif target_or_template == 'template':
            bw_value = Helper.df_template.loc[Helper.df_template['Sequence'] == temp_seq, 'BW' + str(reference_number) + '.50'].iloc[0]  # R132
        else:
            raise Exception('Invalid argument for target or template')

        bw_number = int(bw_value[1:])  # 132
        position_difference = int((reference_number + 0.5 - hotspot_residue_position) * 100)  # 22
        required_position = bw_number - position_difference  # 110
        required_index = required_position - 1  # 109

        # print(hotspot_residue_position, required_position)
        return targ_seq[required_index] if target_or_template == 'target' else temp_seq[required_index]

    # for analysis between sequence chunks
    # see: https://www.ncbi.nlm.nih.gov/blast/html/sub_matrix.html
    @staticmethod
    def identity_for_sequence_chunks(template_seq, target_seq):
        matrix = Helper.gpcr_tm
        for a in pairwise2.align.globaldx(template_seq, target_seq, matrix):
            align_temp = a[0]
            align_targ = a[1]
            seq_len = len(align_temp)
            matches = [align_temp[i] == align_targ[i] for i in range(seq_len)]
            iden = (100 * sum(matches)) / seq_len
            return round(iden, 2)
        return None

    # for analysis between whole sequences
    @staticmethod
    def identity_for_whole_sequence(template_seq, target_seq):
        target_seq = SeqRecord(Seq(target_seq), id="target_seq")
        template_seq = SeqRecord(Seq(template_seq), id="template_seq")

        SeqIO.write(target_seq, "target_seq.fasta", "fasta")
        SeqIO.write(template_seq, "template_seq.fasta", "fasta")

        cline = NcbiblastpCommandline(cmd='blastp', query='target_seq.fasta', subject='template_seq.fasta', outfmt=5)
        stdout_str, stderr_str = cline()
        if stderr_str:
            raise Exception('Error in running Blast command')

        blast_record = NCBIXML.read(StringIO(stdout_str))
        for alignment in blast_record.alignments:
            for hsp in alignment.hsps:
                iden = hsp.identities / hsp.align_length * 100
                return round(iden, 2)

        return None

    @staticmethod
    def align_sequence(template_sequence_chunk, target_sequence_chuck):
        matrix = matlist.blosum62
        return pairwise2.align.globaldx(template_sequence_chunk, target_sequence_chuck, matrix)

    @staticmethod
    def wrap(string, width):
        s = ''
        for i in range(0, len(string), width):
            s += string[i:i + width] + '\n'
        return s

    @staticmethod
    def sanitize_input(sequence):
        import re

        sequence = re.split('(?:>.*\\n)?(.*)', sequence)  # remove '>' and other characters from input sequence (optional), returns list
        sequence = ''.join(sequence)  # join remaining sequence into string
        return re.sub('\\s', '', sequence)  # remove space characters from string

    @staticmethod
    def validate_input(sequence):
        import re

        # check if sequence is empty
        if not sequence:
            return False

        # verify protein sequence
        for residue in re.findall('[a-zA-Z]', sequence):
            if residue not in Helper.valid_residues:
                return False

        return True

    @staticmethod
    def get_unitprot_id(sequence):
        df1 = pd.read_excel(Helper.data_filename, sheet_name="Targets")
        df2 = df1.apply(lambda row: row.astype(str).str.contains(sequence).any(), axis=1)  # searching entered sequence in target database

        if len(df2[df2 == True].index) == 0:
            return False

        df3, = df2[df2 == True].index  # get row number of target sequence
        return df1.loc[df3, "Uniprot_ID"]  # get uniprot id of target sequence

    @staticmethod
    def helical_wheel(sequence, ax, colorcoding='rainbow', lineweights=True, filename=None, seq=False, moment=False, title=''):
        """
        ref: https://www.modlamp.org/modlamp.html#modlamp.plot.helical_wheel
        A function to project a given peptide sequence onto a helical wheel plot. It can be useful to illustrate the
        properties of alpha-helices, like positioning of charged and hydrophobic residues along the sequence.

        :param sequence: {str} the peptide sequence for which the helical wheel should be drawn
        :param colorcoding: {str} the color coding to be used, available: *rainbow*, *charge*, *polar*, *simple*,
            *amphipathic*, *none*
        :param lineweights: {boolean} defines whether connection lines decrease in thickness along the sequence
        :param filename: {str} filename  where to safe the plot. *default = None* --> show the plot
        :param seq: {bool} whether the amino acid sequence should be plotted as a title
        :param moment: {bool} whether the Eisenberg hydrophobic moment should be calculated and plotted
        :param title: {str} title of template or target sequene
        :return: a helical wheel projection plot of the given sequence (interactively or in **filename**)
        :Example:

        >>> helical_wheel('GLFDIVKKVVGALG')
        >>> helical_wheel('KLLKLLKKLLKLLK', colorcoding='charge')
        >>> helical_wheel('AKLWLKAGRGFGRG', colorcoding='none', lineweights=False)
        >>> helical_wheel('ACDEFGHIKLMNPQRSTVWY')

        .. image:: ../docs/static/wheel1.png
            :height: 300px
        .. image:: ../docs/static/wheel2.png
            :height: 300px
        .. image:: ../docs/static/wheel3.png
            :height: 300px
        .. image:: ../docs/static/wheel4.png
            :height: 300px

        .. versionadded:: v2.1.5
        """
        # color mappings
        aa = ['A', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'K', 'L', 'M', 'N', 'P', 'Q', 'R', 'S', 'T', 'V', 'W', 'Y']
        f_rainbow = ['#3e3e28', '#ffcc33', '#b30047', '#b30047', '#ffcc33', '#3e3e28', '#80d4ff', '#ffcc33', '#0047b3',
                     '#ffcc33', '#ffcc33', '#b366ff', '#29a329', '#b366ff', '#0047b3', '#ff66cc', '#ff66cc', '#ffcc33',
                     '#ffcc33', '#ffcc33']
        f_charge = ['#000000', '#000000', '#ff4d94', '#ff4d94', '#000000', '#000000', '#80d4ff', '#000000', '#80d4ff',
                    '#000000', '#000000', '#000000', '#000000', '#000000', '#80d4ff', '#000000', '#000000', '#000000',
                    '#000000', '#000000']
        f_polar = ['#000000', '#000000', '#80d4ff', '#80d4ff', '#000000', '#000000', '#80d4ff', '#000000', '#80d4ff',
                   '#000000', '#000000', '#80d4ff', '#000000', '#80d4ff', '#80d4ff', '#80d4ff', '#80d4ff', '#000000',
                   '#000000', '#000000']
        f_simple = ['#ffcc33', '#ffcc33', '#0047b3', '#0047b3', '#ffcc33', '#7f7f7f', '#0047b3', '#ffcc33', '#0047b3',
                    '#ffcc33', '#ffcc33', '#0047b3', '#ffcc33', '#0047b3', '#0047b3', '#0047b3', '#0047b3', '#ffcc33',
                    '#ffcc33', '#ffcc33']
        f_none = ['#ffffff'] * 20
        f_amphi = ['#ffcc33', '#29a329', '#b30047', '#b30047', '#f79318', '#80d4ff', '#0047b3', '#ffcc33', '#0047b3',
                   '#ffcc33', '#ffcc33', '#80d4ff', '#29a329', '#80d4ff', '#0047b3', '#80d4ff', '#80d4ff', '#ffcc33',
                   '#f79318', '#f79318']
        t_rainbow = ['w', 'k', 'w', 'w', 'k', 'w', 'k', 'k', 'w', 'k', 'k', 'k', 'k', 'k', 'w', 'k', 'k', 'k', 'k', 'k']
        t_charge = ['w', 'w', 'k', 'k', 'w', 'w', 'k', 'w', 'k', 'w', 'w', 'w', 'w', 'w', 'k', 'w', 'w', 'w', 'w', 'w']
        t_polar = ['w', 'w', 'k', 'k', 'w', 'w', 'k', 'w', 'k', 'w', 'w', 'k', 'w', 'k', 'k', 'k', 'k', 'w', 'w', 'w']
        t_simple = ['k', 'k', 'w', 'w', 'k', 'w', 'w', 'k', 'w', 'k', 'k', 'k', 'k', 'w', 'w', 'w', 'w', 'k', 'k', 'k']
        t_none = ['k'] * 20
        t_amphi = ['k', 'k', 'w', 'w', 'w', 'k', 'w', 'k', 'w', 'k', 'k', 'k', 'w', 'k', 'w', 'k', 'k', 'k', 'w', 'w']
        d_eisberg = load_scale('eisenberg')[1]  # eisenberg hydrophobicity values for HM

        if lineweights:
            lw = np.arange(0.1, 5.5, 5. / (len(sequence) - 1))  # line thickness array
            lw = lw[::-1]  # inverse order
        else:
            lw = [2.] * (len(sequence) - 1)

        # check which color coding to use
        if colorcoding == 'rainbow':
            df = dict(zip(aa, f_rainbow))
            dt = dict(zip(aa, t_rainbow))
        elif colorcoding == 'charge':
            df = dict(zip(aa, f_charge))
            dt = dict(zip(aa, t_charge))
        elif colorcoding == 'polar':
            df = dict(zip(aa, f_polar))
            dt = dict(zip(aa, t_polar))
        elif colorcoding == 'simple':
            df = dict(zip(aa, f_simple))
            dt = dict(zip(aa, t_simple))
        elif colorcoding == 'none':
            df = dict(zip(aa, f_none))
            dt = dict(zip(aa, t_none))
        elif colorcoding == 'amphipathic':
            df = dict(zip(aa, f_amphi))
            dt = dict(zip(aa, t_amphi))
        else:
            print("Unknown color coding, 'rainbow' used instead")
            df = dict(zip(aa, f_rainbow))
            dt = dict(zip(aa, t_rainbow))

        # degree to radian
        deg = np.arange(float(len(sequence))) * -100.
        deg = [d + 90. for d in deg]  # start at 270 degree in unit circle (on top)
        rad = np.radians(deg)

        # dict for coordinates and eisenberg values
        d_hydro = dict(zip(rad, [0.] * len(rad)))

        # create figure
        # fig = plt.figure(frameon=False, figsize=(4, 4))
        # ax = fig.add_subplot(111)
        old = None
        hm = list()

        # iterate over sequence
        for i, r in enumerate(rad):
            new = (np.cos(r), np.sin(r))  # new AA coordinates
            if i < 18:
                # plot the connecting lines
                if old is not None:
                    line = lines.Line2D((old[0], new[0]), (old[1], new[1]), transform=ax.transData, color='k', linewidth=lw[i - 1])
                    line.set_zorder(1)  # 1 = level behind circles
                    ax.add_line(line)
            elif 17 < i < 36:
                line = lines.Line2D((old[0], new[0]), (old[1], new[1]), transform=ax.transData, color='k', linewidth=lw[i - 1])
                line.set_zorder(1)  # 1 = level behind circles
                ax.add_line(line)
                new = (np.cos(r) * 1.2, np.sin(r) * 1.2)
            elif i == 36:
                line = lines.Line2D((old[0], new[0]), (old[1], new[1]), transform=ax.transData, color='k', linewidth=lw[i - 1])
                line.set_zorder(1)  # 1 = level behind circles
                ax.add_line(line)
                new = (np.cos(r) * 1.4, np.sin(r) * 1.4)
            else:
                new = (np.cos(r) * 1.4, np.sin(r) * 1.4)

            # plot circles
            circ = patches.Circle(new, radius=0.1, transform=ax.transData, edgecolor='k', facecolor=df[sequence[i]])
            circ.set_zorder(2)  # level in front of lines
            ax.add_patch(circ)

            # check if N- or C-terminus and add subscript, then plot AA letter
            if i == 0:
                ax.text(new[0], new[1], sequence[i] + '$_N$', va='center', ha='center', transform=ax.transData, size=10, color=dt[sequence[i]], fontweight='bold')
            elif i == len(sequence) - 1:
                ax.text(new[0], new[1], sequence[i] + '$_C$', va='center', ha='center', transform=ax.transData, size=10, color=dt[sequence[i]], fontweight='bold')
            else:
                ax.text(new[0], new[1], sequence[i], va='center', ha='center', transform=ax.transData, size=10, color=dt[sequence[i]], fontweight='bold')

            eb = d_eisberg[sequence[i]][0]  # eisenberg value for this AA
            hm.append([eb * new[0], eb * new[1]])  # save eisenberg hydrophobicity vector value to later calculate HM

            old = (np.cos(r), np.sin(r))  # save as previous coordinates

        # draw hydrophobic moment arrow if moment option
        if moment:
            v_hm = np.sum(np.array(hm), 0)
            x = .0333 * v_hm[0]
            y = .0333 * v_hm[1]
            ax.arrow(0., 0., x, y, head_width=0.04, head_length=0.03, transform=ax.transData, color='k', linewidth=6.)
            desc = PeptideDescriptor(sequence)  # calculate hydrophobic moment
            desc.calculate_profile(window=11)
            desc.calculate_moment()
            if abs(x) < 0.2 and y > 0.:  # right positioning of HM text so arrow does not cover it
                z = -0.2
            else:
                z = 0.2

            ax.text(0., z, str(round(desc.descriptor[0][0], 3)), fontdict={'fontsize': 10, 'fontweight': 'bold', 'ha': 'center'})

        # plot shape
        if len(sequence) < 19:
            ax.set_xlim(-1.2, 1.2)
            ax.set_ylim(-1.2, 1.2)
        else:
            ax.set_xlim(-1.4, 1.4)
            ax.set_ylim(-1.4, 1.4)
        ax.spines['right'].set_visible(False)
        ax.spines['top'].set_visible(False)
        ax.spines['left'].set_visible(False)
        ax.spines['bottom'].set_visible(False)
        cur_axes = plt.gca()
        cur_axes.axes.get_xaxis().set_visible(False)
        cur_axes.axes.get_yaxis().set_visible(False)
        plt.tight_layout()

        if seq:
            plt.title(sequence, fontweight='bold', fontsize=10)

        if title:
            ax.text(0., 1.4, title, fontdict={'fontsize': 10, 'fontweight': 'bold', 'ha': 'center'})

        # show or save plot
        if filename:
            plt.savefig(filename, dpi=150)
        # else:
        #     plt.show()

    @staticmethod
    def download_full_alignment(tm_align, template_receptor, target_receptor, template_sequence, target_sequence, template_number, target_number):
        # for template sequence
        nterm_template = template_sequence[:template_number[0][0]]
        icl_1_template = template_sequence[template_number[0][1] + 1:template_number[1][0]]
        ecl_1_template = template_sequence[template_number[1][1] + 1:template_number[2][0]]
        icl_2_template = template_sequence[template_number[2][1] + 1:template_number[3][0]]
        ecl_2_template = template_sequence[template_number[3][1] + 1:template_number[4][0]]
        icl_3_template = template_sequence[template_number[4][1] + 1:template_number[5][0]]
        ecl_3_template = template_sequence[template_number[5][1] + 1:template_number[6][0]]
        cterm_template = template_sequence[template_number[6][1] + 1:]

        # for target sequence
        nterm_target = target_sequence[:target_number[0][0]]
        icl_1_target = target_sequence[target_number[0][1] + 1:target_number[1][0]]
        ecl_1_target = target_sequence[target_number[1][1] + 1:target_number[2][0]]
        icl_2_target = target_sequence[target_number[2][1] + 1:target_number[3][0]]
        ecl_2_target = target_sequence[target_number[3][1] + 1:target_number[4][0]]
        icl_3_target = target_sequence[target_number[4][1] + 1:target_number[5][0]]
        ecl_3_target = target_sequence[target_number[5][1] + 1:target_number[6][0]]
        cterm_target = target_sequence[target_number[6][1] + 1:]

        # align template and target sequences
        nterm_aligned_sequences = Helper.align_sequence(nterm_template, nterm_target)[0]
        icl_1_aligned_sequences = Helper.align_sequence(icl_1_template, icl_1_target)[0]
        ecl_1_aligned_sequences = Helper.align_sequence(ecl_1_template, ecl_1_target)[0]
        icl_2_aligned_sequences = Helper.align_sequence(icl_2_template, icl_2_target)[0]
        ecl_2_aligned_sequences = Helper.align_sequence(ecl_2_template, ecl_2_target)[0]
        icl_3_aligned_sequences = Helper.align_sequence(icl_3_template, icl_3_target)[0]
        ecl_3_aligned_sequences = Helper.align_sequence(ecl_3_template, ecl_3_target)[0]
        cterm_aligned_sequences = Helper.align_sequence(cterm_template, cterm_target)[0]

        # concatenate chunks into whole sequence
        sequence_A = nterm_aligned_sequences.seqA + tm_align[0][0] + \
                     icl_1_aligned_sequences.seqA + tm_align[0][1] + \
                     ecl_1_aligned_sequences.seqA + tm_align[0][2] + \
                     icl_2_aligned_sequences.seqA + tm_align[0][3] + \
                     ecl_2_aligned_sequences.seqA + tm_align[0][4] + \
                     icl_3_aligned_sequences.seqA + tm_align[0][5] + \
                     ecl_3_aligned_sequences.seqA + tm_align[0][6] + \
                     cterm_aligned_sequences.seqA
        sequence_B = nterm_aligned_sequences.seqB + tm_align[1][0] + \
                     icl_1_aligned_sequences.seqB + tm_align[1][1] + \
                     ecl_1_aligned_sequences.seqB + tm_align[1][2] + \
                     icl_2_aligned_sequences.seqB + tm_align[1][3] + \
                     ecl_2_aligned_sequences.seqB + tm_align[1][4] + \
                     icl_3_aligned_sequences.seqB + tm_align[1][5] + \
                     ecl_3_aligned_sequences.seqB + tm_align[1][6] + \
                     cterm_aligned_sequences.seqB

        # wrap string at 80 chars
        wrapped_sequence_A = Helper.wrap(sequence_A, 80)
        wrapped_sequence_B = Helper.wrap(sequence_B, 80)

        return '>' + template_receptor + '\n' + wrapped_sequence_A + \
               '>' + target_receptor + '\n' + wrapped_sequence_B

    @staticmethod
    def calculate_alignment(targ_id, current_item_text):
        df4 = Helper.df_target.apply(lambda row: row.astype(str).str.contains(targ_id).any(), axis=1)  # searching entered sequence in target database
        df5, = df4[df4 == True].index  # getting row number of target sequence
        target_id = Helper.df_target.loc[df5, 'Uniprot_ID']
        target_receptor = Helper.df_target.loc[df5, 'Receptor']
        target_sequence = Helper.df_target.loc[df5, 'Sequence']

        template_number = []
        target_number = []

        for i in range(1, 8):
            target_number.append((Helper.df_target.loc[df5, f'TM{i} start'], Helper.df_target.loc[df5, f'TM{i} end']))

        df2 = Helper.df_template.apply(lambda row: row.astype(str).str.contains(current_item_text).any(), axis=1)  # searching clicked name in template database
        df3, = df2[df2 == True].index  # getting row number of template name
        template_id = Helper.df_template.loc[df3, 'PDBID']  # getting pdbid of template name
        template_receptor = Helper.df_template.loc[df3, 'Receptor']
        template_sequence = Helper.df_template.loc[df3, 'Sequence']

        for j in range(1, 8):
            template_number.append((Helper.df_template.loc[df3, f'TM{j} start'], Helper.df_template.loc[df3, f'TM{j} end']))

        return (
            target_id,
            target_receptor,
            target_sequence,
            target_number,
            template_id,
            template_receptor,
            template_sequence,
            template_number,
        )

    @staticmethod
    def show_hydrophobicity_plot(ress, hydrh, summ, value, template_receptor, target_receptor, name, download=False):
        if len(ress) == 0:
            QtWidgets.QMessageBox.about(QtWidgets.QMessageBox(), "Information", "Please select a template first")
            return

        hydrophobicity_plot_filenames = []
        os.makedirs('images', exist_ok=True)

        for i in range(7):
            plt.figure('TM ' + str(i + 1))

            line = [0, 0]
            colors = ['r', 'b']

            for j in range(2):
                line[j], = plt.plot(ress[i], hydrh[i][j], color=colors[j], linewidth=1.0, linestyle="-")

            plt.xlabel('Residue Number')
            plt.ylabel('Mean Hydrophobicity Value')

            if template_receptor != '':
                template_receptor =  template_receptor
            if target_receptor != '':
                target_receptor =  target_receptor

            plt.legend([line[0], line[1]], ['Template: ' + template_receptor, 'Target: ' + target_receptor])
            plt.title('Helix ' + str(i + 1) + ': SSD per residue=%.4f' % (value[i]))

            if download:
                filename = Helper.images_dir + '/' + name + '_helix_' + str(i + 1) + '.png'
                hydrophobicity_plot_filenames.append(filename)
                plt.savefig(filename)
                plt.cla()
                plt.close()

        # download plot
        if download:
            return hydrophobicity_plot_filenames

        plt.show()
        plt.cla()
        plt.close()

    @staticmethod
    def show_helical_plot(template, target, template_receptor, target_receptor, name, download=False):
        helical_plot_filenames = []
        os.makedirs('images', exist_ok=True)

        for i in range(7):
            fig, (ax1, ax2) = plt.subplots(1, 2)  # 1 row, 2 col
            fig.set_size_inches((10, 5))  # size of plot in inches 10x5
            fig.canvas.set_window_title('TM ' + str(i + 1))

            if template_receptor != '':
                template_receptor = ': ' + template_receptor
            if target_receptor != '':
                target_receptor = ': ' + target_receptor

            Helper.helical_wheel(template[0][i], ax1, colorcoding='amphipathic', moment=True, title='Template' + template_receptor)
            Helper.helical_wheel(target[0][i], ax2, colorcoding='amphipathic', moment=True, title='Target' + target_receptor)

            # remove x and y axes
            ax1.set_axis_off()
            ax2.set_axis_off()
            fig.add_axes(ax1)
            fig.add_axes(ax2)

            if download:
                filename = Helper.images_dir + '/' + name + '_helical_' + str(i + 1) + '.png'
                helical_plot_filenames.append(filename)
                plt.savefig(filename)
                plt.cla()
                plt.close(fig)

        # download plot
        if download:
            return helical_plot_filenames

        plt.show()
        plt.cla()
        plt.close()

    @staticmethod
    def generate_result_summary(tm_align, template_receptor, target_receptor, template_sequence, target_sequence, template_number, target_number, hydrophobicity_plots, helical_plots):
        document = Document()
        styles = document.styles

        # Courier, 8pt
        fixed_width_style = styles.add_style('FixedWidthStyle', WD_STYLE_TYPE.CHARACTER)
        fixed_width_style.base_style = styles['Normal']
        fixed_width_style.font.name = 'Courier'
        fixed_width_style.font.size = Pt(8)

        # document content goes here...

        # page #1
        # heading
        document.add_heading(f'Output from {template_receptor}-{target_receptor}', level=1)
        # citation
        document.add_paragraph('Please cite: Jabeen A., Ranganathan, S. BIO-GATS: A tool for automated GPCR template selection through a biophysical approach. (unpublished)', style='Intense Quote')
        # full alignment
        alignment_text = Helper.download_full_alignment(tm_align, template_receptor, target_receptor, template_sequence, target_sequence, template_number, target_number)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(alignment_text)
        run.style = fixed_width_style

        # page #2-8
        for i in range(7):
            # page break
            document.add_page_break()
            # heading
            document.add_heading(f'TM{i + 1}', level=2)
            # template/target alignment
            paragraph = document.add_paragraph()
            run = paragraph.add_run('Template: ' + tm_align[0][i] + '\n' + 'Target:   ' + tm_align[1][i])
            run.style = fixed_width_style
            # hydrophobicity plot
            document.add_picture(hydrophobicity_plots[i], width=Inches(6.))
            # helical wheel plot
            document.add_picture(helical_plots[i], width=Inches(6.))

        try:
            document.save(Helper.summary_filename)
            QtWidgets.QMessageBox.about(QtWidgets.QMessageBox(), "Information", "Results summary downloaded in result-summary.docx")
        except PermissionError:
            QtWidgets.QMessageBox.about(QtWidgets.QMessageBox(), "Error", "File result-summary.docx already in use. Close and try again.")
        finally:
            for i in range(7):
                os.remove(hydrophobicity_plots[i])
                os.remove(helical_plots[i])
